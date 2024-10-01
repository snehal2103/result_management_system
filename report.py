from tkinter import *
from PIL import Image, ImageTk
from PIL import Image, ImageDraw, ImageFont
import tempfile
import subprocess
from docx import Document
from tkinter import filedialog
from tkinter import ttk, messagebox
import os
import sqlite3
class reportClass:
    def __init__(self, root):
        self.root = root
        self.root.title("Student Result Management System")
        self.root.geometry("1200x480+80+170")
        self.root.config(bg="white")
        self.root.focus_force()
        self.print_window = None
        title = Label(self.root, text="View Student Result", font=("goudy old style", 20, "bold"), bg="orange", fg="#262626")
        title.place(x=10, y=15, width=1180, height=50)
        self.var_search = StringVar()
        self.var_id = ""
        lbl_search = Label(self.root, text="Search By Roll No.:", font=("goudy old style", 20, 'bold'), bg='white')
        lbl_search.place(x=280, y=100)
        entry_search = Entry(self.root, textvariable=self.var_search, font=("goudy old style", 20), bg='lightblue')
        entry_search.place(x=520, y=100, width=150)
        btn_search = Button(self.root, text="Search", font=("goudy old style", 15, "bold"), bg="blue", fg="white", cursor="hand2", command=self.search)
        btn_search.place(x=680, y=100, width=100, height=35)
        btn_clear = Button(self.root, text="Clear", font=("goudy old style", 15, "bold"), bg="grey", fg="white", cursor="hand2", command=self.clear)
        btn_clear.place(x=800, y=100, width=100, height=35)
        labels = ["Roll No:", "Name:", "Course:", "Marks Obtained:", "Total Marks:", "Percentage:"]
        for i in range(len(labels)):
            lbl = Label(self.root, text=labels[i], font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)
            lbl.place(x=150 + i * 150, y=230, width=150, height=50)
        self.roll = Label(self.root, font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)
        self.roll.place(x=150, y=280, width=150, height=50)
        self.name = Label(self.root, font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)
        self.name.place(x=300, y=280, width=150, height=50)
        self.course = Label(self.root, font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)
        self.course.place(x=450, y=280, width=150, height=50)
        self.marks = Label(self.root, font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)
        self.marks.place(x=600, y=280, width=150, height=50)
        self.full = Label(self.root, font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)
        self.full.place(x=750, y=280, width=150, height=50)
        self.per = Label(self.root, font=("goudy old style", 15, 'bold'), bg='white', bd=2, relief=GROOVE)

        self.per.place(x=900, y=280, width=150, height=50)
       
        btn_delete = Button(self.root, text="Delete", font=("goudy old style", 15, "bold"), bg="red", fg="white", cursor="hand2", command=self.delete)
        btn_delete.place(x=450, y=350, width=150, height=35)
        
        btn_print = Button(self.root, text="Print", font=("goudy old style", 15, "bold"), bg="green", fg="white", cursor="hand2", command=self.print_result)
        btn_print.place(x=620, y=350, width=150, height=35)
        
    def search(self):
        con = sqlite3.connect(database="rms.db")
        cur = con.cursor()
        try:
            if self.var_search.get() == "":
                messagebox.showerror("Error", "Roll no. is required", parent=self.root)
            else:
                cur.execute("select * from result where roll=?", (self.var_search.get(),))
                row = cur.fetchone()
                if row != None:
                    self.var_id = row[0]
                    self.roll.config(text=row[1])
                    self.name.config(text=row[2])
                    self.course.config(text=row[3])
                    self.marks.config(text=row[4])
                    self.full.config(text=row[5])
                    percentage = float(row[6])
                    pass_status = "Pass" if percentage >= 40 else "Fail"
                    self.per.config(text=f"{percentage} ({pass_status}%)")
                else:
                    messagebox.showerror("Error", "No record found", parent=self.root)
        except Exception as ex:
            messagebox.showerror("Error", f"Error due to {str(ex)}")
    
    # ... (previous code remains unchanged)

    def print_result(self):
        if self.var_id == "":
            messagebox.showerror("Error", "Search Student result first", parent=self.root)
        else:
            # Open a separate window for printing
            self.print_window = Toplevel(self.root)
            self.print_window.title("Print Result")
            self.print_window.geometry("400x200+400+200")
            self.print_window.grab_set()  # Make the window modal
            print_label = Label(self.print_window, text="Printing is in progress, please wait...", font=("goudy old style", 15))
            print_label.pack(pady=50)
            # Call the method to handle printing
            self.print_result_details()
    def print_result_details(self):
        try:
            # Get student and result details from the database based on the roll number
            con = sqlite3.connect(database="rms.db")
            cur = con.cursor()
            cur.execute("SELECT * FROM result WHERE roll=?", (self.var_search.get(),))
            result_details = cur.fetchone()
            cur.execute("SELECT * FROM student WHERE roll=?", (self.var_search.get(),))
            student_details = cur.fetchone()
            con.close()

            if student_details and result_details:
                # Create a Word document with student and result details
                doc = Document()
                doc.add_heading('Student and Result Details', level=1)

                # Add student details to the document
                doc.add_heading('Student Details', level=2)
                doc.add_paragraph(f'Name: {student_details[1]}')
                doc.add_paragraph(f'Roll No: {student_details[0]}')
                doc.add_paragraph(f'DOB: {student_details[4]}')
                doc.add_paragraph(f'Email: {student_details[2]}')
                doc.add_paragraph(f'Contact No: {student_details[5]}')
                doc.add_paragraph(f'Admission Date: {student_details[6]}')
                doc.add_paragraph(f'Course: {student_details[7]}')

                # Add result details to the document
                doc.add_heading('Result Details', level=2)
                doc.add_paragraph(f'Marks Obtained: {result_details[4]}')
                doc.add_paragraph(f'Total Marks: {result_details[5]}')
                doc.add_paragraph(f'Percentage: {result_details[6]}')
                percentage = float(result_details[6])
                status = "Pass" if percentage >= 40 else "Fail"
                doc.add_paragraph(f'Status: {status}')
                performance = self.get_performance(percentage)
                doc.add_paragraph(f'Performance: {performance}')

                # Save the student and result details to a temporary file
                file_path = filedialog.asksaveasfilename(parent=self.print_window, defaultextension=".docx", filetypes=[("Word files", "*.docx")])
                if file_path:
                    doc.save(file_path)
                    messagebox.showinfo("Success", "Student and result details exported to Word document successfully!", parent=self.print_window)
                else:
                    messagebox.showwarning("Warning", "Printing cancelled", parent=self.print_window)
                self.print_window.destroy()  # Close the printing window after completion or cancellation

            else:
                messagebox.showerror("Error", "No student or result record found", parent=self.print_window)
                self.print_window.destroy()  # Close the printing window if there is an error

        except Exception as ex:
            messagebox.showerror("Error", f"Error: {str(ex)}", parent=self.print_window)
            self.print_window.destroy()  # Close the printing window if there is an error


    def get_performance(self, percentage):
     if percentage >= 85:
        return "Excellent"
     elif percentage >= 70:
        return "Good"
     elif percentage >= 50:
        return "Satisfactory"
     else:
        return "Needs Improvement"

# Method to handle print button in the View Result window
    def print_result_from_view(self):
        self.print_result()
# Method to open View Result window from the Dashboard without closing it after printing
    def open_view_result_window(self):
        if self.view_result_window is None or not self.view_result_window.winfo_exists():
            self.view_result_window = Toplevel(self.root)
            self.view_result_window.geometry("600x400")  # Set the dimensions as per your requirement

            # Populate the View Result window with student and result details here
            # ...

            # Add a print button in the View Result window
            print_button = Button(self.view_result_window, text="Print", font=("goudy old style", 15, "bold"), bg="green", fg="white", cursor="hand2", command=self.print_result)
            print_button.place(x=250, y=300, width=100, height=35)
    def clear(self):
        self.var_id = ""
        self.roll.config(text="")
        self.name.config(text="")
        self.course.config(text="")
        self.marks.config(text="")
        self.full.config(text="")
        self.per.config(text="")
        self.var_search.set("")
  
    def delete(self):
        con = sqlite3.connect(database="rms.db")
        cur = con.cursor()
        try:
            if self.var_id == "":
                messagebox.showerror("Error", "Search Student result first", parent=self.root)
            else:
                cur.execute("select * from result where rid=?", (self.var_id,))
                row = cur.fetchone()
                if row == None:
                    messagebox.showerror("Error", "Invalid Student Result", parent=self.root)
                else:
                    op = messagebox.askyesno("Confirm", "Do you really want to delete?", parent=self.root)
                    if op == True:
                        cur.execute("delete from result where rid=?", (self.var_id,))
                        con.commit()
                        messagebox.showinfo("Delete", "Result deleted Successfully", parent=self.root)
                        self.clear()
        except Exception as ex:
            messagebox.showerror("Error", f"Error due to {str(ex)}")

if __name__ == "__main__":
    root = Tk()
    obj = reportClass(root)
    root.mainloop()